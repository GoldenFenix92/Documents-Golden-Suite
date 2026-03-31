from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml import OxmlElement, ns
import os

class APAFormatter:
    def __init__(self, file_path, version, fuente_seleccionada, paragraphs_data, style_map_inverse, nuevas_referencias=None):
        self.file_path = file_path
        self.version = version
        self.paragraphs_data = paragraphs_data
        self.style_map_inverse = style_map_inverse
        self.nuevas_referencias = nuevas_referencias if nuevas_referencias else []
        self.doc = Document(file_path)
        
        partes_fuente = fuente_seleccionada.rsplit(" ", 1)
        self.font_name = partes_fuente[0]
        self.font_size = int(partes_fuente[1])

    def crear_numero_pagina(self, run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def aplicar_reglas_globales(self):
        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)
            
            header = section.header
            for p in header.paragraphs:
                p.text = ""
            
            p_header = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_header = p_header.add_run()
            run_header.font.name = self.font_name
            run_header.font.size = Pt(self.font_size)
            self.crear_numero_pagina(run_header)

        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)
        
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def formatear_bloque(self, paragraph, estilo_espanol):
        pf = paragraph.paragraph_format
        pf.first_line_indent = 0
        pf.left_indent = 0
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        if not paragraph.runs:
            paragraph.add_run(paragraph.text)
        
        for run in paragraph.runs:
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            run.bold = False
            run.italic = False

            if estilo_espanol == "Párrafo Normal":
                pf.first_line_indent = Cm(1.27)
            
            elif estilo_espanol == "Referencia":
                pf.left_indent = Cm(1.27)
                pf.first_line_indent = Cm(-1.27)
                
            elif estilo_espanol == "Cita en Bloque":
                pf.left_indent = Cm(1.27)
                pf.first_line_indent = 0
            
            elif estilo_espanol in ["Título Principal", "Título 1"]:
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
            
            elif estilo_espanol == "Título 2":
                run.bold = True
            
            elif estilo_espanol == "Título 3":
                if "7a" in self.version:
                    run.bold = True
                    run.italic = True
                else: 
                    pf.first_line_indent = Cm(1.27)
                    run.bold = True

    def procesar(self, solo_indices=None):
        self.aplicar_reglas_globales()
        
        # Procesar parrafos existentes
        for p_data in self.paragraphs_data:
            idx = p_data["id"]
            if solo_indices is not None and idx not in solo_indices:
                continue
                
            real_paragraph = self.doc.paragraphs[idx]
            self.formatear_bloque(real_paragraph, p_data["current_style"])

        # Inyectar nuevas referencias al final del documento
        if self.nuevas_referencias:
            # Añadir salto de pagina para la seccion de bibliografia
            self.doc.add_page_break()
            
            # Añadir titulo de referencias
            p_titulo = self.doc.add_paragraph("Referencias")
            self.formatear_bloque(p_titulo, "Título 1")
            
            # Añadir cada referencia con sangria francesa
            for ref in self.nuevas_referencias:
                p_ref = self.doc.add_paragraph(ref)
                self.formatear_bloque(p_ref, "Referencia")

        nombre_base, ext = os.path.splitext(self.file_path)
        ruta_salida = f"{nombre_base}_APA_PROCESADO{ext}"
        self.doc.save(ruta_salida)
        return ruta_salida